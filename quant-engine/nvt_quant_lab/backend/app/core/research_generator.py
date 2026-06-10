import os
import json
import io
import pandas as pd
import numpy as np
from datetime import datetime
from typing import List, Dict, Any, Tuple, Optional

# Try imports for python-docx
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from core.ai_advisor import _get_model

DISCLAIMER_TEXT = "This report is generated using quantitative outputs and AI-assisted interpretation. It is for research and educational purposes only and does not constitute investment advice."

def generate_template_research(
    analysis_type: str,
    quant: Optional[dict],
    backtest: Optional[dict],
    optimizer: Optional[dict],
    benchmark: str,
    lang: str
) -> dict:
    """
    Fallback deterministic template-based generator when Gemini is not configured.
    Outputs high-quality financial summary utilizing actual figures from the engines.
    """
    is_vi = (lang == "vi")
    
    # Extract metrics safely
    symbols = []
    if quant and "tickers" in quant:
        symbols = quant["tickers"]
    elif backtest and "strategy" in backtest and "symbols" in backtest["strategy"]:
        symbols = backtest["strategy"]["symbols"]
    elif optimizer and "weights" in optimizer:
        symbols = list(optimizer["weights"].keys())

    # Capital
    capital = 100000000.0
    if backtest and "strategy" in backtest and "initial_capital" in backtest["strategy"]:
        capital = backtest["strategy"]["initial_capital"]
    elif optimizer and "metrics" in optimizer and "initial_capital" in optimizer: # or config
        pass

    # Expected Return / Volatility
    q_ret = quant.get("metrics", {}).get("expected_return") if quant else None
    q_vol = quant.get("metrics", {}).get("volatility") if quant else None
    q_sharpe = quant.get("metrics", {}).get("sharpe_ratio") if quant else None
    q_mdd = quant.get("metrics", {}).get("max_drawdown") if quant else None

    # Backtest Metrics
    bt_ret = backtest.get("metrics", {}).get("annualized_return") if backtest else None
    bt_mdd = backtest.get("metrics", {}).get("max_drawdown") if backtest else None
    bt_sharpe = backtest.get("metrics", {}).get("sharpe_ratio") if backtest else None
    bt_costs = backtest.get("costs", {}).get("total_costs", 0) if backtest else 0

    # Optimizer Metrics
    opt_ret = optimizer.get("metrics", {}).get("expected_return") if optimizer else None
    opt_vol = optimizer.get("metrics", {}).get("volatility") if optimizer else None
    opt_sharpe = optimizer.get("metrics", {}).get("sharpe_ratio") if optimizer else None
    opt_method = optimizer.get("optimizer") if optimizer else None

    symbols_str = ", ".join(symbols)

    if is_vi:
        exec_summary = f"Báo cáo phân tích đầu tư cho danh mục gồm các mã: {symbols_str}. Báo cáo được xây dựng dựa trên dữ liệu lịch sử và các giải pháp định lượng nâng cao tại NVT Quant Lab. "
        if optimizer:
            exec_summary += f"Danh mục được tối ưu hóa bằng phương pháp '{opt_method}' với mục tiêu cân đối tỷ suất lợi nhuận kỳ vọng và rủi ro biến động."
        else:
            exec_summary += "Danh mục được đánh giá theo phương pháp phân bổ tỷ trọng đều baseline."

        perf_analysis = "Không có đủ dữ liệu phân tích hiệu suất."
        if q_ret is not None or bt_ret is not None or opt_ret is not None:
            perf_analysis = f"Lợi nhuận kỳ vọng của danh mục đạt khoảng {fmt_pct(opt_ret or q_ret or bt_ret)}. "
            if bt_ret is not None:
                perf_analysis += f"Trong quá trình kiểm thử lịch sử (backtest), tỷ suất lợi nhuận năm đạt {fmt_pct(bt_ret)}. "
            if opt_ret is not None and q_ret is not None:
                perf_analysis += f"Hiệu suất sau tối ưu hóa ({fmt_pct(opt_ret)}) cải thiện so với danh mục baseline phân bổ đều ({fmt_pct(q_ret)})."

        risk_analysis = "Độ biến động kỳ vọng và tỷ số Sharpe cho thấy mức độ chấp nhận rủi ro ở mức trung bình. "
        vol_val = opt_vol or q_vol
        if vol_val is not None:
            risk_analysis += f"Độ biến động kỳ vọng (Volatility) được kiểm soát ở mức {fmt_pct(vol_val)}. "
        mdd_val = bt_mdd or q_mdd
        if mdd_val is not None:
            risk_analysis += f"Mức sụt giảm tài sản lớn nhất lịch sử (Max Drawdown) là {fmt_pct(mdd_val)}. "
        sharpe_val = opt_sharpe or q_sharpe or bt_sharpe
        if sharpe_val is not None:
            risk_analysis += f"Tỷ số Sharpe Ratio đạt {sharpe_val:.2f}, biểu thị mức bù đắp lợi nhuận hợp lý trên mỗi đơn vị rủi ro chịu đựng."

        bench_analysis = f"Danh mục được đánh giá hiệu suất so với chỉ số tham chiếu {benchmark}. "
        if quant and quant.get("metrics", {}).get("beta") is not None:
            beta = quant["metrics"]["beta"]
            bench_analysis += f"Hệ số Beta của danh mục so với thị trường là {beta:.2f}, thể hiện độ nhạy tương đối với các biến động của chỉ số {benchmark}."

        obs = f"Vốn đầu tư ban đầu ước tính: {capital:,.0f} VND. "
        if backtest:
            obs += f"Chi phí giao dịch phát sinh và trượt giá trong backtest là {bt_costs:,.0f} VND. Việc tái cân bằng thường xuyên cần cân nhắc chi phí ma sát này."

        takeaways = [
            f"Danh mục tập trung vào nhóm cổ phiếu: {symbols_str}.",
            "Tối ưu hóa tỷ trọng giúp cải thiện hiệu quả sử dụng vốn.",
            "Cần theo dõi sát chỉ số sụt giảm tài sản lớn nhất để quản trị rủi ro đuôi.",
            "Kiểm soát tần suất tái cơ cấu để hạn chế chi phí ma sát giao dịch."
        ]
    else:
        exec_summary = f"Investment analysis report for portfolio: {symbols_str}. The report is constructed using historical datasets and advanced quantitative engines at NVT Quant Lab. "
        if optimizer:
            exec_summary += f"The portfolio is optimized using the '{opt_method}' method to balance expected return and volatility."
        else:
            exec_summary += "The portfolio is analyzed using the baseline equal-weight allocation."

        perf_analysis = "Insufficient data for performance analysis."
        if q_ret is not None or bt_ret is not None or opt_ret is not None:
            perf_analysis = f"The expected annualized return of the portfolio is approximately {fmt_pct(opt_ret or q_ret or bt_ret)}. "
            if bt_ret is not None:
                perf_analysis += f"During the historical backtest, the annualized return achieved was {fmt_pct(bt_ret)}. "
            if opt_ret is not None and q_ret is not None:
                perf_analysis += f"The optimized return ({fmt_pct(opt_ret)}) shows enhancement compared to the equal-weight baseline ({fmt_pct(q_ret)})."

        risk_analysis = "Expected volatility and Sharpe ratios indicate moderate risk tolerance. "
        vol_val = opt_vol or q_vol
        if vol_val is not None:
            risk_analysis += f"Annualized expected volatility is managed at {fmt_pct(vol_val)}. "
        mdd_val = bt_mdd or q_mdd
        if mdd_val is not None:
            risk_analysis += f"The worst historical peak-to-trough drop (Max Drawdown) was observed at {fmt_pct(mdd_val)}. "
        sharpe_val = opt_sharpe or q_sharpe or bt_sharpe
        if sharpe_val is not None:
            risk_analysis += f"The Sharpe Ratio stands at {sharpe_val:.2f}, indicating reasonable excess return per unit of risk."

        bench_analysis = f"The portfolio performance is compared against the benchmark {benchmark}. "
        if quant and quant.get("metrics", {}).get("beta") is not None:
            beta = quant["metrics"]["beta"]
            bench_analysis += f"The portfolio Beta is {beta:.2f}, describing its relative sensitivity compared to the {benchmark} index."

        obs = f"Initial capital allocated: {capital:,.0f} VND. "
        if backtest:
            obs += f"Total transaction and slippage costs incurred during the simulation amounted to {bt_costs:,.0f} VND. Friction costs should be monitored closely."

        takeaways = [
            f"Portfolio focuses on equities: {symbols_str}.",
            "Optimal weight allocation increases portfolio efficiency.",
            "Drawdowns must be monitored closely to hedge tail risk.",
            "Rebalancing frequencies should balance weights accuracy and friction costs."
        ]

    return {
        "executive_summary": exec_summary,
        "performance_analysis": perf_analysis,
        "risk_analysis": risk_analysis,
        "benchmark_analysis": bench_analysis,
        "portfolio_observations": obs,
        "key_takeaways": takeaways
    }

def fmt_pct(val: Optional[float]) -> str:
    if val is None:
        return "--"
    return f"{val * 100:.2f}%"

def generate_grounded_research(
    analysis_type: str,
    quant: Optional[dict],
    backtest: Optional[dict],
    optimizer: Optional[dict],
    benchmark: str = "VN30",
    language: str = "vi"
) -> Tuple[dict, List[str]]:
    warnings = []
    
    # 1. Verification of inputs
    if not quant:
        warnings.append("Cảnh báo: Dữ liệu phân tích định lượng (quant_results) bị thiếu. Báo cáo sử dụng dữ liệu mặc định.")
    if not backtest:
        warnings.append("Cảnh báo: Dữ liệu kiểm thử lịch sử (backtest_results) bị thiếu. Báo cáo không chứa kiểm thử chi tiết.")
    if not optimizer:
        warnings.append("Cảnh báo: Dữ liệu tối ưu hóa (optimizer_results) bị thiếu. Sử dụng phân bổ đều làm mặc định.")

    # Get Gemini model
    model = _get_model()
    
    if model is None:
        warnings.append("Gemini API Key chưa được cấu hình hoặc bị lỗi. Đang trả về báo cáo theo biểu mẫu template định sẵn.")
        fallback = generate_template_research(analysis_type, quant, backtest, optimizer, benchmark, language)
        return fallback, warnings

    # 2. Build structured prompt for grounding
    data_payload = {
        "analysis_type": analysis_type,
        "quant_results": quant or {},
        "backtest_results": backtest or {},
        "optimizer_results": optimizer or {},
        "benchmark": benchmark,
        "language": language
    }
    
    prompt = f"""
You are a Principal AI Investment Analyst at NVT Quant Lab.
Construct a highly professional investment research report of type '{analysis_type}' based strictly on the quantitative inputs.

INPUT DATA (JSON):
{json.dumps(data_payload, indent=2)}

INSTRUCTIONS:
1. Ground every claim, number, return, volatility, Sharpe ratio, and drawdown figure directly in the provided JSON data. 
2. DO NOT make up, invent, or extrapolate any numbers. If a metrics field is missing, describe it as not available.
3. Write your output in language: '{language}' (use professional Vietnamese if 'vi', professional English if 'en').
4. The output must be returned strictly in JSON format. Do not add markdown or backticks (no ```json).
5. The JSON structure must contain exactly these keys:
   - executive_summary (string): High-level executive overview.
   - performance_analysis (string): Analyze returns, CAGR, and capital growth.
   - risk_analysis (string): Analyze volatility, Sharpe, Sortino, drawdowns, and risk contributions.
   - benchmark_analysis (string): Compare portfolio against the benchmark index.
   - portfolio_observations (string): Observations about allocations, rebalancing, and trading costs.
   - key_takeaways (array of strings): 4 specific investment takeaways.

JSON Response:
"""

    try:
        response = model.generate_content(prompt)
        text = response.text.strip()
        
        # Clean markdown code blocks if Gemini returns them
        if text.startswith("```"):
            lines = text.split("\n")
            if lines[0].startswith("```json") or lines[0].startswith("```"):
                lines = lines[1:]
            if lines[-1].startswith("```"):
                lines = lines[:-1]
            text = "\n".join(lines).strip()
            
        research = json.loads(text)
        
        # Verify required keys are present
        required_keys = ["executive_summary", "performance_analysis", "risk_analysis", "benchmark_analysis", "portfolio_observations", "key_takeaways"]
        for k in required_keys:
            if k not in research:
                research[k] = ""
        if not isinstance(research["key_takeaways"], list):
            research["key_takeaways"] = [str(research["key_takeaways"])]
            
        return research, warnings
        
    except Exception as e:
        warnings.append(f"Lỗi gọi Gemini API ({str(e)}). Đang trả về báo cáo mẫu template fallback.")
        fallback = generate_template_research(analysis_type, quant, backtest, optimizer, benchmark, language)
        return fallback, warnings

def export_research_to_docx(research_data: dict, language: str = "vi") -> bytes:
    """
    Generates a professionally-formatted Word Document from the research data.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed or available on this system.")

    doc = Document()
    is_vi = (language == "vi")
    
    # Configure styles
    title_text = "BÁO CÁO PHÂN TÍCH ĐẦU TƯ AI" if is_vi else "AI INVESTMENT RESEARCH REPORT"
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = 'Arial'
    
    # Subtitle with date
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    sub_run = sub.add_run(f"NVT Quant Lab Research | Date: {date_str}")
    sub_run.font.size = Pt(10)
    sub_run.italic = True
    
    # Add a horizontal line
    doc.add_paragraph("__________________________________________________________________").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sections = [
        ("1. Executive Summary", research_data.get("executive_summary", "")),
        ("2. Portfolio Composition", 
         ("Cấu trúc và tỷ trọng phân bổ tài sản định lượng." if is_vi else "Quantitative asset allocation structure.")),
        ("3. Quantitative Performance Review", research_data.get("performance_analysis", "")),
        ("4. Backtest Findings", 
         (f"Kết quả mô phỏng kiểm thử lịch sử chi tiết. {research_data.get('performance_analysis', '')}" if is_vi 
          else f"Detailed historical backtesting insights. {research_data.get('performance_analysis', '')}")),
        ("5. Optimizer Findings", 
         (f"Kết quả tối ưu hóa theo thuyết danh mục hiện đại. {research_data.get('portfolio_observations', '')}" if is_vi 
          else f"Optimization results based on Modern Portfolio Theory. {research_data.get('portfolio_observations', '')}")),
        ("6. Benchmark Comparison", research_data.get("benchmark_analysis", "")),
        ("7. Risk Assessment", research_data.get("risk_analysis", "")),
        ("8. Portfolio Strengths", 
         ("Thuật toán tối ưu hóa giúp giảm thiểu biến động và tối đa hóa hiệu suất sử dụng vốn." if is_vi 
          else "Optimal weight diversification maximizes Sharpe ratio and capital efficiency.")),
        ("9. Portfolio Weaknesses", 
         ("Hiệu suất phụ thuộc vào các cú sốc thanh khoản vĩ mô và rủi ro thị trường hệ thống." if is_vi 
          else "Performance remains sensitive to macro liquidity and systemic shocks.")),
        ("10. Important Limitations", 
         ("Mô hình được xây dựng từ chuỗi dữ liệu lịch sử. Kết quả backtest không cam kết hiệu suất thực tế tương lai." if is_vi 
          else "Backtest results are historical simulations and do not guarantee future performance.")),
        ("11. Conclusion", 
         ("Danh mục đầu tư được thiết kế khoa học giúp cân đối rủi ro-lợi nhuận tối ưu." if is_vi 
          else "The quantitative allocation strategy successfully balances risk and return objectives."))
    ]
    
    for label, content in sections:
        # Add heading
        h = doc.add_paragraph()
        hrun = h.add_run(label)
        hrun.font.size = Pt(13)
        hrun.bold = True
        hrun.font.name = 'Arial'
        
        # Add content
        doc.add_paragraph(content)
        
    # Key takeaways
    h_takeaways = doc.add_paragraph()
    tk_run = h_takeaways.add_run("Khuyến nghị then chốt / Key Takeaways")
    tk_run.font.size = Pt(13)
    tk_run.bold = True
    tk_run.font.name = 'Arial'
    
    for item in research_data.get("key_takeaways", []):
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.add_run(item)
        
    # Disclaimer
    doc.add_paragraph() # Spacer
    p_disc = doc.add_paragraph()
    d_run = p_disc.add_run(DISCLAIMER_TEXT)
    d_run.font.size = Pt(9)
    d_run.italic = True
    d_run.font.name = 'Arial'
    
    # Save to binary stream
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream.getvalue()
